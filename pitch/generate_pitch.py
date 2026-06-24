# -*- coding: utf-8 -*-
"""
EVENT AI — генератор инвест-питч пакета
Выход: EVENT-AI-Pitch-RU.docx (python-docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Цвета
ACCENT_BLUE = RGBColor(0x2C, 0x7B, 0xFF)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)
LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)

OUTPUT_PATH = r"D:\event-ai\pitch\EVENT-AI-Pitch-RU.docx"


def set_heading_color(paragraph, color):
    for run in paragraph.runs:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    """Добавляет заголовок с синим цветом"""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = ACCENT_BLUE
        if level == 1:
            run.font.size = Pt(18)
            run.bold = True
        else:
            run.font.size = Pt(14)
            run.bold = True
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    """Добавляет обычный параграф"""
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = DARK_GRAY
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Добавляет буллет-пункт"""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(11)
        run_bold.font.color.rgb = DARK_GRAY
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(11)
        run_rest.font.color.rgb = DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
    return p


def add_numbered(doc, text, bold_prefix=None):
    """Добавляет нумерованный пункт"""
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(11)
        run_bold.font.color.rgb = DARK_GRAY
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(11)
        run_rest.font.color.rgb = DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
    return p


def add_divider(doc):
    """Добавляет горизонтальную линию"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2C7BFF")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_placeholder(doc, text):
    """Добавляет плейсхолдер (выделенный жирным серым)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)  # тёмно-янтарный — визуально отличается
    return p


def build_docx():
    doc = Document()

    # Поля документа
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Стиль Normal — основной шрифт
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    # ────────────────────────────────────────────
    # 1. ТИТУЛ
    # ────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("EVENT AI")
    run_title.bold = True
    run_title.font.size = Pt(32)
    run_title.font.color.rgb = ACCENT_BLUE

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Инвестиционное предложение")
    run_sub.font.size = Pt(18)
    run_sub.bold = True
    run_sub.font.color.rgb = DARK_GRAY

    doc.add_paragraph()  # пустая строка

    p_tagline = doc.add_paragraph()
    p_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tag = p_tagline.add_run(
        "AI-маркетплейс организации мероприятий «под ключ» — первый в Казахстане."
    )
    run_tag.italic = True
    run_tag.font.size = Pt(13)
    run_tag.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("Разработчик: студия DEADLINE  |  Контакт: +7 747 253 11 38")
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = LIGHT_GRAY

    add_divider(doc)
    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 2. КРАТКОЕ ОПИСАНИЕ (ELEVATOR PITCH)
    # ────────────────────────────────────────────
    add_heading(doc, "1. О проекте", level=1)
    add_para(
        doc,
        "EVENT AI — первый в Казахстане AI-маркетплейс для организации мероприятий. "
        "Пользователь описывает своё событие в чате с AI — и мгновенно получает полный пакет: "
        "подборку проверенных подрядчиков, смету, сценарий, тайминг и возможность сразу забронировать всё необходимое. "
        "Мы устраняем хаос организации тоев, свадеб и корпоративов, объединяя весь рынок подрядчиков в одном месте "
        "и превращая многонедельный процесс — в несколько кликов."
    )

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 3. ПРОБЛЕМА
    # ────────────────────────────────────────────
    add_heading(doc, "2. Проблема", level=1)
    add_para(doc, "Организация мероприятия сегодня — это стресс, время и деньги:")
    bullets_problem = [
        "Площадку, ведущих, артистов, декор, фото- и видеографов приходится искать по отдельности — в Instagram, по рекомендациям, в чатах.",
        "Нет прозрачной сметы: цены скрыты, сравнить предложения сложно — клиент переплачивает или получает не то, что ожидал.",
        "Подрядчики не верифицированы — высокий риск нарваться на непрофессионалов или потерять предоплату.",
        "Координация всех участников ложится на плечи клиента: тайминг, логистика, резервные планы — всё вручную.",
        "В Казахстане нет единой платформы, которая объединяла бы клиентов и поставщиков event-услуг — рынок фрагментирован."
    ]
    for b in bullets_problem:
        add_bullet(doc, b)

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 4. РЕШЕНИЕ
    # ────────────────────────────────────────────
    add_heading(doc, "3. Решение — что делает EVENT AI", level=1)
    add_para(
        doc,
        "EVENT AI — это AI-ассистент и маркетплейс в одном приложении. "
        "Клиент описывает мероприятие текстом или голосом, а система автоматически собирает всё необходимое:"
    )
    bullets_solution = [
        ("Подборка подрядчиков — ", "верифицированные площадки, ведущие, артисты, декораторы, фотографы и видеографы по вашему бюджету и формату."),
        ("Прозрачная смета — ", "AI рассчитывает стоимость мероприятия с разбивкой по статьям, без скрытых платежей."),
        ("Готовый сценарий и тайминг — ", "структура события с тайм-кодами, которую можно редактировать и передать подрядчикам."),
        ("Пригласительные — ", "генерация дизайна и рассылка гостям прямо из приложения."),
        ("Бронирование онлайн — ", "оплата бронирования подрядчиков без звонков и переписки.")
    ]
    for bp, bt in bullets_solution:
        add_bullet(doc, bt, bold_prefix=bp)

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 5. КАК РАБОТАЕТ
    # ────────────────────────────────────────────
    add_heading(doc, "4. Как это работает", level=1)
    steps = [
        ("Клиент открывает приложение и описывает событие — ", "формат, дату, количество гостей, бюджет."),
        ("AI анализирует запрос — ", "подбирает подходящих подрядчиков из базы платформы."),
        ("Клиент видит готовую подборку — ", "с ценами, рейтингами, портфолио, отзывами."),
        ("AI генерирует смету, сценарий и тайминг — ", "одним нажатием."),
        ("Бронирование и оплата — ", "прямо в приложении; подрядчик получает заявку мгновенно.")
    ]
    for sp, st in steps:
        add_numbered(doc, st, bold_prefix=sp)

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 6. РЫНОК И ПОЧЕМУ СЕЙЧАС
    # ────────────────────────────────────────────
    add_heading(doc, "5. Рынок и почему сейчас", level=1)
    add_para(
        doc,
        "Казахстан — перспективный рынок с богатой культурой мероприятий. "
        "Тои, свадьбы, корпоративы, дни рождения — массовый, регулярный, высокобюджетный сегмент."
    )
    market_bullets = [
        "Алматы — крупнейший мегаполис страны; event-рынок активен круглый год.",
        "Высокая средняя стоимость мероприятия — тои и свадьбы в Казахстане традиционно масштабны.",
        "Цифровизация услуг набирает скорость: потребитель привык заказывать онлайн (Kaspi, Glovo, 2GIS).",
        "Нет ни одного AI-маркетплейса для event-подрядчиков — ниша полностью свободна.",
        "Первопроходец занимает рынок быстро: сеть подрядчиков и доверие клиентов — барьер для конкурентов.",
        "Модель легко масштабируется на другие города Казахстана и рынки СНГ."
    ]
    for b in market_bullets:
        add_bullet(doc, b)

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 7. БИЗНЕС-МОДЕЛЬ
    # ────────────────────────────────────────────
    add_heading(doc, "6. Бизнес-модель", level=1)
    biz_bullets = [
        ("Комиссия 7% — ", "с каждого подтверждённого заказа через платформу. Основной поток выручки."),
        ("Бронирование — ", "50 000 ₸ за факт подтверждённого бронирования — клиент платит сразу, безопасная транзакция."),
        ("Платное продвижение подрядчиков — ", "приоритетное размещение в подборках и каталоге платформы."),
        ("Премиум-размещение — ", "топовые позиции для крупных площадок и агентств."),
        ("B2B-пакеты — ", "тарифы для event-компаний, площадок и корпоративных заказчиков: расширенный профиль, CRM-интеграция, аналитика.")
    ]
    for bp, bt in biz_bullets:
        add_bullet(doc, bt, bold_prefix=bp)

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 8. ЧТО УЖЕ ГОТОВО
    # ────────────────────────────────────────────
    add_heading(doc, "7. Что уже готово", level=1)
    add_para(doc, "Мы показываем инвестору готовую упаковку, а не просто идею:")
    ready_bullets = [
        ("Рабочее демо-приложение — ", "живая ссылка: https://deadline-corp.github.io/event-ai-demo/"),
        ("Инвесторская презентация — ", "6 слайдов с ключевыми показателями и концепцией."),
        ("Анимационный видеоролик — ", "40+ секунд, демонстрирует продукт и ценностное предложение."),
        ("ТЗ на разработку — ", "детальное техническое задание на полноценное приложение."),
        ("Мобильное приложение — ", "в разработке (App Store / Google Play).")
    ]
    for bp, bt in ready_bullets:
        add_bullet(doc, bt, bold_prefix=bp)

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 9. ПРЕИМУЩЕСТВА ПРОЕКТА
    # ────────────────────────────────────────────
    add_heading(doc, "8. Почему EVENT AI победит", level=1)
    adv_bullets = [
        "Большой, понятный рынок — тои и свадьбы в Казахстане не исчезнут; спрос стабилен и предсказуем.",
        "Первый игрок — нет прямых конкурентов с AI-функционалом на казахстанском рынке.",
        "Простая модель монетизации — комиссия с транзакций; выручка растёт вместе с объёмом.",
        "AI снижает ручную работу — меньше менеджеров нужно для обработки заявок; маржа выше.",
        "Двусторонняя сеть — чем больше подрядчиков, тем лучше подборки; чем больше клиентов, тем выгоднее подрядчикам; сеть самоусиливается.",
        "Масштабируемость — Алматы → Астана → другие города Казахстана → рынки СНГ без изменения бизнес-модели.",
        "Готовность к запуску — демо, ТЗ, видео, презентация уже готовы; инвестиции ускорят, а не начнут проект."
    ]
    for b in adv_bullets:
        add_bullet(doc, b)

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 10. ИНВЕСТИЦИОННОЕ ПРЕДЛОЖЕНИЕ (ПЛЕЙСХОЛДЕР)
    # ────────────────────────────────────────────
    add_heading(doc, "9. Инвестиционное предложение", level=1)
    add_para(
        doc,
        "Мы ищем партнёра-инвестора, который разделяет видение и готов войти на раннем этапе — "
        "когда рынок ещё свободен и оценка минимальна."
    )
    doc.add_paragraph()
    add_placeholder(
        doc,
        "▶  [Сумма инвестиций и условия — заполняется после расчёта сметы]"
    )
    add_placeholder(
        doc,
        "▶  [Доля / структура сделки — заполняется после переговоров]"
    )
    add_placeholder(
        doc,
        "▶  [Pre-money оценка проекта — заполняется после финансового моделирования]"
    )
    add_placeholder(
        doc,
        "▶  [На что идут инвестиции: разработка / маркетинг / команда — детализировать]"
    )
    add_placeholder(
        doc,
        "▶  [Ожидаемый возврат / выход для инвестора — заполняется]"
    )

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 11. УСТНЫЙ ПИТЧ (60–90 сек)
    # ────────────────────────────────────────────
    add_divider(doc)
    add_heading(doc, "10. Устный питч (60–90 секунд)", level=1)
    add_para(
        doc,
        "Скрипт для живой встречи с инвестором — читается вслух за 60–90 секунд:",
        bold=True
    )
    doc.add_paragraph()

    pitch_text = (
        "Представьте: вам нужно организовать той на 200 человек. Вы открываете десятки вкладок в Instagram, "
        "пишете в WhatsApp незнакомым людям, собираете цены вручную — и через неделю всё ещё не знаете, "
        "в кого вложить предоплату. Это не исключение — это норма для казахстанского рынка.\n\n"
        "EVENT AI решает эту боль одним приложением. Клиент описывает своё мероприятие в AI-чате — "
        "и за несколько секунд получает подборку проверенных подрядчиков, прозрачную смету, готовый сценарий "
        "и тайминг, пригласительные — и сразу может забронировать всё необходимое. "
        "Никаких звонков, никаких пропавших предоплат, никакого стресса.\n\n"
        "Мы — первый AI-маркетплейс event-услуг в Казахстане. Ниша свободна. "
        "Рабочее демо уже готово — вы можете зайти и попробовать прямо сейчас. "
        "Монетизация простая: 7% комиссии с каждого подтверждённого заказа. "
        "Чем больше мероприятий — тем больше выручка, без роста затрат.\n\n"
        "Нам нужен партнёр-инвестор, чтобы перейти от демо к полноценному приложению в App Store. "
        "Если вам интересно войти первым в этот рынок — давайте обсудим условия."
    )

    p_pitch = doc.add_paragraph()
    run_pitch = p_pitch.add_run(pitch_text)
    run_pitch.font.size = Pt(11)
    run_pitch.font.color.rgb = DARK_GRAY
    run_pitch.italic = True

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # 12. ПЕРВОЕ СООБЩЕНИЕ ИНВЕСТОРУ
    # ────────────────────────────────────────────
    add_heading(doc, "11. Шаблон первого сообщения инвестору", level=1)
    add_para(
        doc,
        "Шаблон для WhatsApp / Telegram / Email (3–5 предложений):",
        bold=True
    )
    doc.add_paragraph()

    msg_text = (
        "Добрый день! Меня зовут [Ваше имя], я основатель стартапа EVENT AI — AI-маркетплейс для организации мероприятий в Казахстане. "
        "Мы первые в нише: клиент описывает событие в чате с AI и за секунды получает подборку подрядчиков, смету и сценарий. "
        "Рабочее демо уже готово: [ссылка на демо]. "
        "Я ищу инвестора на раннем этапе — рынок свободен, вход сейчас максимально выгоден. "
        "Готов провести короткую встречу или созвон в удобное для вас время — скажите, интересно?"
    )

    p_msg = doc.add_paragraph()
    p_msg.style = doc.styles["Normal"]
    # Рамка вокруг абзаца (визуальный блок)
    pPr2 = p_msg._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "8")
        bdr.set(qn("w:color"), "2C7BFF")
        pBdr2.append(bdr)
    pPr2.append(pBdr2)
    run_msg = p_msg.add_run(msg_text)
    run_msg.font.size = Pt(11)
    run_msg.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    # ────────────────────────────────────────────
    # FOOTER — контакт
    # ────────────────────────────────────────────
    add_divider(doc)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run(
        "EVENT AI  ·  Разработчик: студия DEADLINE  ·  +7 747 253 11 38  ·  https://deadline-corp.github.io/event-ai-demo/"
    )
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = LIGHT_GRAY

    # Сохранение
    doc.save(OUTPUT_PATH)
    print(f"Сохранено: {OUTPUT_PATH}")

    # Верификация — открываем и считаем параграфы
    doc2 = Document(OUTPUT_PATH)
    para_count = len(doc2.paragraphs)
    print(f"Параграфов в документе: {para_count}")

    return para_count


if __name__ == "__main__":
    count = build_docx()
    print(f"✅ Готово. Параграфов: {count}")
